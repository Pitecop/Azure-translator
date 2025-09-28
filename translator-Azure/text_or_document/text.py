import requests
from docx import Document
import os
subscripion_key = "sua chave aqui"
endpoint = 'O seu ponto de extremidade aqui'
location = "sua localizacao aqui"
language_destination = "linguagemd e destino aqui"

def translator_text(text, target_language):
    path = '/translate'
    constructed_url = endpoint + path
    headers = {
        'Ocp-Apim-Subscription-Key': subscripion_key,
        'Ocp-Apim-Subscription-Region': location,
        'Content-type': 'application/json',
        'X-ClientTraceId' : str(os.urandom(16))
        }
    Body = [{
        'text': text
    }]
    params = {
        'api-version' : '3.0',
        'from' : 'en',
        'to': target_language
    }
    request = requests.post(constructed_url, params=params, headers=headers, json=Body)
    response = request.json()
    print("API Response:", response) # Print the full response for debugging
    try:
        return response[0]['translations'][0]['text']
    except KeyError:
        print("Error: Could not parse the API response. Please check the subscription key, endpoint, and other parameters.")
        return None
    
translator_text("One of the girls - The Weekend", language_destination)

#para para leitura de documentos
def translator_document(path):
    document = Document(path)
    full_text = []
    for paragraph in document.paragraphs:
     translated_text = translator_text(paragraph.text, language_destination)
     full_text.append(translated_text)

    translated_doc = Document()
    for line in full_text:
        translated_doc.add_paragraph(line)
    path_translated = path.replace(".docx", f"_{language_destination}.docx")
    translated_doc.save(path_translated)
    return path_translated

input_file = "caminho do seu arquivo aqui"
translator_document(input_file)